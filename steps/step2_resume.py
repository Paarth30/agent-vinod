import os
import shutil
import hashlib
import difflib
from pathlib import Path
from rich.console import Console
from rich.prompt import Confirm, Prompt
from rich.panel import Panel
from rich.table import Table
from rich.text import Text
import anthropic
import config as cfg
from steps.step1_discover import _job_key
from steps.llm import gemini_call

console = Console()

RESUME_RULES = """
ATS & Resume Best Practices (apply all):
1. Use standard section headings: Summary, Experience, Education, Skills, Projects.
2. Start every bullet with a strong action verb (Led, Built, Reduced, Increased, Designed, etc.).
3. Quantify achievements wherever possible (%, $, time saved, users, scale).
4. Mirror exact keywords and phrases from the job description naturally.
5. Remove any skills or experiences irrelevant to this specific role.
6. Keep sentences concise — no paragraphs in bullet points.
7. No tables, columns, text boxes, headers/footers, or images in content (ATS breaks on these).
8. Use consistent date format: Mon YYYY (e.g. Jan 2023).
9. Do not fabricate experience — only enhance/reframe what exists.
10. Keep the summary (2-3 sentences) role-specific and punchy.
"""


def run(jobs: list[dict]) -> list[dict]:
    console.print("\n[bold]Step 2: Resume Tailoring[/bold]")

    resume_docx = Path(cfg.RESUME_DOCX_PATH)
    if not resume_docx.exists():
        console.print(f"[red]Resume not found at {resume_docx}[/red]")
        raise SystemExit(0)

    output_dir = Path("data/resumes")
    output_dir.mkdir(parents=True, exist_ok=True)

    client = anthropic.Anthropic(api_key=cfg.GEMINI_API_KEY)
    original_text = _extract_docx_text(resume_docx)

    # ── Phase 1: tailor every resume back-to-back, no prompts ──────────────
    console.print(f"  Tailoring {len(jobs)} resume(s)...\n")
    for i, job in enumerate(jobs, 1):
        if not job.get("jd"):
            console.print(f"  [{i}/{len(jobs)}] [yellow]{job['company']} — no JD, using original[/yellow]")
            job["_resume_text"] = None
            continue

        tailored_text = _tailor(client, original_text, job)
        if not tailored_text:
            console.print(f"  [{i}/{len(jobs)}] [yellow]{job['company']} — Claude returned empty, using original[/yellow]")
            job["_resume_text"] = None
            continue

        job["_resume_text"] = tailored_text
        before = _ats_score(original_text, job)
        after  = _ats_score(tailored_text, job)
        b, a = before.get("score") or 0, after.get("score") or 0
        delta = a - b
        delta_str = f"[green]+{delta}[/green]" if delta > 0 else (f"[red]{delta}[/red]" if delta < 0 else "[dim]=[/dim]")
        console.print(f"  [{i}/{len(jobs)}] [green]OK[/green] {job['company']} — {job['title']}  ATS {b}%→{a}% ({delta_str})")

    for job in jobs:
        _save_resume(job, resume_docx, output_dir)

    # ── Phase 2: review list — open a specific one to see the diff and chat ─
    _review_loop(jobs, client, original_text, resume_docx, output_dir)

    if not Confirm.ask("\n[bold green]Resumes ready — proceed to cover letters?[/bold green]"):
        console.print("[red]Aborted.[/red]")
        raise SystemExit(0)

    return jobs


def _save_resume(job: dict, resume_docx: Path, output_dir: Path):
    """Write the job's current _resume_text to disk (or fall back to the original)."""
    text = job.get("_resume_text")
    if not text:
        job["resume_docx"] = str(resume_docx)
        job["resume_pdf"] = None
        return

    safe_company = "".join(c if c.isalnum() or c in "-_ " else "" for c in job["company"]).replace(" ", "_")
    safe_title   = "".join(c if c.isalnum() or c in "-_ " else "" for c in job["title"]).replace(" ", "_")
    key_hash = hashlib.md5(_job_key(job).encode()).hexdigest()[:8]
    out_path = output_dir / f"{safe_company}_{safe_title}_{key_hash}.docx"
    _write_docx(resume_docx, text, out_path)
    pdf_path = _convert_to_pdf(out_path)

    job["resume_docx"] = str(out_path)
    job["resume_pdf"] = str(pdf_path) if pdf_path else None


def _review_loop(jobs: list[dict], client, original_text: str, resume_docx: Path, output_dir: Path):
    while True:
        _show_review_table(jobs, original_text)
        console.print(
            "\n  [dim]Enter a job number to review/edit its resume, "
            "or press Enter to accept all and continue.[/dim]"
        )
        choice = Prompt.ask("  Selection", default="").strip()

        if not choice or choice.lower() == "done":
            return
        if not choice.isdigit() or not (1 <= int(choice) <= len(jobs)):
            console.print("[yellow]  Invalid selection.[/yellow]")
            continue

        _review_job(jobs[int(choice) - 1], client, original_text, resume_docx, output_dir)


def _show_review_table(jobs: list[dict], original_text: str):
    t = Table(title="Tailored Resumes", show_lines=True)
    t.add_column("#", style="cyan", width=4)
    t.add_column("Company", style="yellow")
    t.add_column("Title", style="white")
    t.add_column("ATS Before→After")
    t.add_column("Status")

    for i, job in enumerate(jobs, 1):
        text = job.get("_resume_text")
        if text is None:
            status = "[dim]Original (no JD)[/dim]" if not job.get("jd") else "[yellow]Using original[/yellow]"
            t.add_row(str(i), job["company"], job["title"], "[dim]n/a[/dim]", status)
            continue
        before = _ats_score(original_text, job)
        after  = _ats_score(text, job)
        b, a = before.get("score") or 0, after.get("score") or 0
        t.add_row(str(i), job["company"], job["title"], f"{b}% → {a}%", "[green]Tailored[/green]")

    console.print("\n", t)


def _review_job(job: dict, client, original_text: str, resume_docx: Path, output_dir: Path):
    current_text = job.get("_resume_text")
    if current_text is None:
        console.print("[yellow]  No tailored resume for this job — nothing to review.[/yellow]")
        return

    before_ats = _ats_score(original_text, job)
    while True:
        _show_diff(original_text, current_text)
        _show_ats_panel(before_ats, _ats_score(current_text, job), job)
        console.print(
            "\n  [dim]Type feedback to refine (e.g. 'make bullet 3 more concise'), "
            "[cyan]regen[/cyan] to regenerate from scratch, [cyan]skip[/cyan] to use the original, "
            "or press Enter to go back to the list.[/dim]"
        )
        action = Prompt.ask("  Feedback", default="").strip()

        if not action or action.lower() == "done":
            break
        elif action.lower() == "skip":
            current_text = None
            break
        elif action.lower() == "regen":
            console.print("  [dim]Regenerating...[/dim]")
            regenerated = _tailor(client, original_text, job)
            if regenerated:
                current_text = regenerated
            else:
                console.print("[yellow]  Regeneration failed — keeping previous version.[/yellow]")
        else:
            console.print("  [dim]Applying your changes...[/dim]")
            revised = _apply_feedback(client, current_text, action, job)
            if revised:
                current_text = revised
            else:
                console.print("[yellow]  Edit failed — keeping current version.[/yellow]")

    job["_resume_text"] = current_text
    _save_resume(job, resume_docx, output_dir)
    if current_text is None:
        console.print("  [yellow]Using original resume for this job.[/yellow]")
    else:
        console.print(f"  [green]OK[/green] Saved to [dim]{job.get('resume_docx')}[/dim]")


# ── ATS scoring ────────────────────────────────────────────────────────────────

def _ats_score(text: str, job: dict) -> dict:
    from steps.step1_discover import _ats_score as _score
    return _score(text, job)


def _show_ats_panel(before: dict, after: dict, job: dict):
    b_score = before.get("score") or 0
    a_score = after.get("score") or 0
    b_bd    = before.get("breakdown") or {}
    a_bd    = after.get("breakdown") or {}

    LABEL_COLORS = {"Excellent": "green", "Good": "cyan", "Fair": "yellow", "Low": "red"}

    def _score_str(score, label):
        color = LABEL_COLORS.get(label, "white")
        return f"[{color}]{score}% {label}[/{color}]"

    def _delta(b, a):
        d = a - b
        if d > 0:  return f"[green]+{d}[/green]"
        if d < 0:  return f"[red]{d}[/red]"
        return "[dim]=[/dim]"

    t = Table(show_header=True, header_style="bold dim", box=None, padding=(0, 2))
    t.add_column("Category",      style="dim",  min_width=16)
    t.add_column("Before",        justify="right", min_width=12)
    t.add_column("After",         justify="right", min_width=12)
    t.add_column("Change",        justify="right", min_width=8)

    rows = [
        ("Skills /35",     "skills"),
        ("Keywords /30",   "keywords"),
        ("Experience /20", "experience"),
        ("Education /10",  "education"),
        ("Title /5",       "title"),
    ]
    for label, key in rows:
        b, a = b_bd.get(key, 0), a_bd.get(key, 0)
        t.add_row(label, str(b), str(a), _delta(b, a))

    t.add_section()
    t.add_row(
        "[bold]TOTAL[/bold]",
        _score_str(b_score, before.get("label", "")),
        _score_str(a_score, after.get("label", "")),
        f"[bold]{_delta(b_score, a_score)}[/bold]",
    )

    console.print(Panel(t, title=f"[bold]ATS Score — {job['company']}[/bold]", expand=False))

    missing = after.get("missing_skills", [])
    if missing:
        console.print(f"  [dim]Skills still missing: {', '.join(missing[:8])}[/dim]")


# ── Diff display ───────────────────────────────────────────────────────────────

def _show_diff(original: str, revised: str):
    orig_lines = original.splitlines()
    rev_lines  = revised.splitlines()

    diff = list(difflib.unified_diff(orig_lines, rev_lines, lineterm="", n=2))

    if not diff:
        console.print(Panel("[dim]No changes from original.[/dim]", title="Diff", expand=False))
        return

    rich_text = Text()
    for line in diff[2:]:  # skip the --- +++ header lines
        if line.startswith("+"):
            rich_text.append(line + "\n", style="green")
        elif line.startswith("-"):
            rich_text.append(line + "\n", style="red strike")
        elif line.startswith("@@"):
            rich_text.append(line + "\n", style="cyan dim")
        else:
            rich_text.append(line + "\n", style="dim")

    console.print(Panel(rich_text, title="[bold]Original  [red]-removed[/red]  [green]+added[/green][/bold]", expand=False))


# ── Claude calls ───────────────────────────────────────────────────────────────

_FORMAT_RULES = """
STRICT OUTPUT FORMAT — follow exactly, no exceptions:
• Line 1:   Candidate's full name only (no title, no degree)
• Line 2:   Contact details, pipe-separated — e.g.  +91 98765 43210  |  email@example.com  |  linkedin.com/in/handle
• Line 3:   Leave blank
• Sections: Section name in ALL CAPS on its own line, no colon — e.g. SUMMARY, EXPERIENCE, EDUCATION, SKILLS
• Job entries: Company Name  |  Job Title  |  Mon YYYY – Mon YYYY  (on one line, immediately after the section header)
• Bullets:  Each achievement on its own line, starting with "• " (bullet + space)
• No markdown, no asterisks, no bold markers, no extra symbols, no tables, no columns.
• Leave one blank line between sections.
"""


def _tailor(client, original_text: str, job: dict) -> str:
    prompt = f"""You are an expert resume writer and ATS specialist.

Here is the candidate's current resume:
<resume>
{original_text}
</resume>

Here is the job description they are applying to:
<job_title>{job['title']}</job_title>
<company>{job['company']}</company>
<job_description>
{job['jd'][:3000]}
</job_description>

Resume writing rules:
{RESUME_RULES}

{_FORMAT_RULES}

Rewrite the resume tailored to this specific job, following the format rules above exactly.
Output ONLY the resume content — no preamble, no explanation, nothing else.
"""
    return gemini_call(client, prompt, max_tokens=4096)


def _apply_feedback(client, current_text: str, feedback: str, job: dict) -> str:
    prompt = f"""You are editing a resume for a {job['title']} role at {job['company']}.

Current resume:
<resume>
{current_text}
</resume>

Candidate's feedback:
<feedback>
{feedback}
</feedback>

{_FORMAT_RULES}

Apply the feedback precisely. Do not change anything not mentioned in the feedback.
Preserve the exact format rules above — ALL CAPS section headers, "• " bullets, pipe-separated job lines.
Output ONLY the full updated resume. No commentary, no explanation.
"""
    return gemini_call(client, prompt, max_tokens=4096)




# ── Docx helpers ───────────────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _write_docx(template_path: Path, new_text: str, out_path: Path):
    """
    Build a clean, ATS-friendly Word document from Claude's plain-text output.
    Page margins are copied from the original template; all content is rebuilt
    with professional formatting (bold name, section borders, bullet indents, etc.).
    """
    import re
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shutil.copy(template_path, out_path)
    doc = Document(str(out_path))

    # Clear all body content; preserve the trailing sectPr (page margins / paper size)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sdt"):
            body.remove(child)

    # ── Style helpers ──────────────────────────────────────────────────────────

    NAVY   = RGBColor(0x1F, 0x39, 0x64)
    GREY   = RGBColor(0x55, 0x55, 0x55)
    DGREY  = RGBColor(0x33, 0x33, 0x33)

    def _r(p, text, bold=False, italic=False, size=10, color=None):
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
        return run

    def _border(p):
        """Thin navy bottom rule under section headers."""
        pPr   = p._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "1F3964")
        pBdr.append(bot)
        pPr.append(pBdr)

    SECTIONS = {
        "SUMMARY", "PROFESSIONAL SUMMARY", "OBJECTIVE", "PROFILE",
        "EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE",
        "EDUCATION", "ACADEMIC BACKGROUND",
        "SKILLS", "TECHNICAL SKILLS", "KEY SKILLS", "CORE COMPETENCIES",
        "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS", "AWARDS",
        "PUBLICATIONS", "VOLUNTEERING", "LANGUAGES", "INTERESTS", "REFERENCES",
    }

    def _classify(text, name_done, contact_done):
        upper = text.upper().rstrip(":")
        if not name_done:
            return "name"
        if not contact_done and re.search(
            r"@|linkedin\.com|github\.com|\+?\d[\d\s()\-]{6,}", text, re.I
        ):
            return "contact"
        if upper in SECTIONS:
            return "section"
        # Catch ALL-CAPS short lines that are likely section headers
        alpha = [c for c in text if c.isalpha()]
        if alpha and 3 < len(text) < 40 and sum(1 for c in alpha if c.isupper()) / len(alpha) > 0.8:
            return "section"
        if re.match(r"^[•·▪\-]\s", text) or text.startswith("• "):
            return "bullet"
        return "body"

    # ── Write lines ────────────────────────────────────────────────────────────

    name_done    = False
    contact_done = False

    for raw in new_text.split("\n"):
        text = raw.strip()
        if not text:
            continue

        kind = _classify(text, name_done, contact_done)

        if kind == "name":
            p = doc.add_paragraph()
            _r(p, text, bold=True, size=20, color=NAVY)
            p.paragraph_format.space_after = Pt(1)
            name_done = True

        elif kind == "contact":
            p = doc.add_paragraph()
            _r(p, text, size=9, color=GREY)
            p.paragraph_format.space_after = Pt(8)
            contact_done = True

        elif kind == "section":
            p = doc.add_paragraph()
            _r(p, text.upper().rstrip(":"), bold=True, size=10, color=NAVY)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            _border(p)

        elif kind == "bullet":
            content = re.sub(r"^[•·▪\-]\s*", "", text).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent        = Inches(0.25)
            p.paragraph_format.first_line_indent  = Inches(-0.15)
            p.paragraph_format.space_after        = Pt(1)
            _r(p, "•  " + content, size=9.5)

        else:  # body — job title lines, education lines, free text
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            # "Company  |  Role  |  Date" → bold company, normal rest
            parts = re.split(r"\s*\|\s*", text, maxsplit=2)
            if len(parts) >= 2:
                _r(p, parts[0].strip(), bold=True, size=9.5)
                for part in parts[1:]:
                    _r(p, "  |  ", size=9.5, color=GREY)
                    _r(p, part.strip(), size=9.5, color=DGREY)
            else:
                _r(p, text, size=9.5)

    doc.save(str(out_path))


def _convert_to_pdf(docx_path: Path):
    try:
        from docx2pdf import convert
        pdf_path = docx_path.with_suffix(".pdf")
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception as e:
        console.print(f"[yellow]  PDF conversion failed (Word not installed?): {e}[/yellow]")
        return None
